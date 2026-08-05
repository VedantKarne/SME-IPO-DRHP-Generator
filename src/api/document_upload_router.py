"""
src/api/document_upload_router.py

FastAPI router for client document uploads.
Provides two endpoints:
  POST /api/documents/upload/{company_id}   — Upload a file, trigger background processing
  GET  /api/documents/status/{company_id}   — Poll per-document processing status

Background job pipeline:
  1. Save file to Original_Docs/Client/{company_id}/
  2. Route to correct parser (pdf_parser for PDFs, openpyxl for Excel, python-docx for Word)
  3. For each page → client_data_chunker.process() → classify to DRHP section
  4. Embed chunks → index to Vector Store 'client_documents' collection
  5. Auto-populate SQLite tables: financial_statement, uploaded_document
"""
import os
import uuid
import logging
import shutil
from typing import List, Optional

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks, Depends
from sqlalchemy.orm import Session
from pydantic import BaseModel

from src.api.auth_router import get_current_user
from src.extraction.schema import Company, UploadedDocument, FinancialStatement, AuditLog
from src.extraction.db_session import SessionLocal
from src.ingestion.pdf_parser import parse_pdf
from src.ingestion.client_data_chunker import ClientDataChunker
from src.retrieval.vector_store import VectorStore, COLLECTION_CLIENT

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/documents", tags=["documents"])

CLIENT_DOCS_DIR = "Uploaded_Docs"
ACCEPTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".doc"}


# ─── DB Dependency ────────────────────────────────────────────────────────────

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# ─── Response Models ──────────────────────────────────────────────────────────

class UploadResponse(BaseModel):
    upload_id: str
    filename: str
    status: str
    message: str


class DocumentStatusItem(BaseModel):
    upload_id: str
    filename: str
    doc_type: str
    status: str
    error_message: Optional[str] = None
    uploaded_at: str


# ─── Helper: Extract Text from Various File Types ─────────────────────────────

def _extract_text_from_file(file_path: str) -> str:
    """Routes extraction to the correct parser based on file extension."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        parsed_docs = parse_pdf(file_path, source="client")
        return "\n".join(d.text for d in parsed_docs if d.text)

    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet in wb.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(c) if c is not None else "" for c in row)
                    if row_text.strip():
                        lines.append(row_text)
            return "\n".join(lines)
        except ImportError:
            logger.warning("openpyxl not installed. Cannot parse Excel files.")
            return ""

    elif ext in (".docx", ".doc"):
        try:
            import docx
            document = docx.Document(file_path)
            return "\n".join(p.text for p in document.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed. Cannot parse Word files.")
            return ""

    logger.warning(f"Unsupported file extension: {ext}")
    return ""


# ─── Helper: Auto-extract Financial KPIs and populate SQLite ──────────────────

def _try_auto_populate_financials(
    text: str, company_id_uuid: uuid.UUID, filename: str, db: Session
):
    """
    Attempts a simple keyword + Groq extraction to populate the financial_statement
    table with high-level KPIs (Revenue, PAT, Net Worth, EBITDA, etc.).
    Only runs for documents classified as financial statements.
    """
    try:
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            return

        from groq import Groq
        import json as json_lib

        client = Groq(api_key=api_key)
        prompt = (
            "You are a financial data extractor. From the text below, extract annual financial metrics "
            "for each available fiscal year. Return ONLY a JSON array (no markdown) like:\n"
            '[{"fiscal_year": 2024, "revenue_lakhs": 6200, "ebitda_lakhs": 1150, '
            '"pat_lakhs": 800, "net_worth_lakhs": 3700, "paid_up_capital_lakhs": 500}]\n\n'
            "If a value is not found, use null. Only return the JSON array.\n\n"
            f"Text:\n{text[:8000]}"
        )
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=512,
        )
        raw = completion.choices[0].message.content.strip()
        # Strip potential markdown code fences
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        records = json_lib.loads(raw)

        for rec in records:
            fy = rec.get("fiscal_year")
            if not fy:
                continue
            existing = (
                db.query(FinancialStatement)
                .filter(
                    FinancialStatement.company_id == company_id_uuid,
                    FinancialStatement.fiscal_year == fy,
                )
                .first()
            )
            if existing:
                # Update only if AI provides non-null values
                existing.revenue_lakhs = rec.get("revenue_lakhs") or existing.revenue_lakhs
                existing.ebitda_lakhs = rec.get("ebitda_lakhs") or existing.ebitda_lakhs
                existing.pat_lakhs = rec.get("pat_lakhs") or existing.pat_lakhs
                existing.net_worth_lakhs = rec.get("net_worth_lakhs") or existing.net_worth_lakhs
                existing.paid_up_capital_lakhs = rec.get("paid_up_capital_lakhs") or existing.paid_up_capital_lakhs
                existing.source = "ai_extracted"
            else:
                db.add(
                    FinancialStatement(
                        company_id=company_id_uuid,
                        fiscal_year=fy,
                        revenue_lakhs=rec.get("revenue_lakhs"),
                        ebitda_lakhs=rec.get("ebitda_lakhs"),
                        pat_lakhs=rec.get("pat_lakhs"),
                        net_worth_lakhs=rec.get("net_worth_lakhs"),
                        paid_up_capital_lakhs=rec.get("paid_up_capital_lakhs"),
                        source="ai_extracted",
                    )
                )
        db.commit()
        logger.info(f"Auto-populated {len(records)} financial year records from '{filename}'")
    except Exception as e:
        logger.warning(f"Could not auto-populate financials from '{filename}': {e}")


# ─── Background Processing Job ────────────────────────────────────────────────

def _process_document_background(
    file_path: str,
    upload_id: str,
    company_id_str: str,
    company_name: str,
    filename: str,
):
    """
    Background task: parse → chunk → embed → index → update status.
    Runs asynchronously after the API returns 202.
    """
    db = SessionLocal()
    try:
        company_id_uuid = uuid.UUID(company_id_str)

        # Mark as processing
        record = db.query(UploadedDocument).filter(UploadedDocument.id == uuid.UUID(upload_id)).first()
        if record:
            record.status = "processing"
            db.commit()

        # 1. Extract text
        logger.info(f"Extracting text from '{filename}'...")
        text = _extract_text_from_file(file_path)

        if not text.strip():
            raise ValueError("No text could be extracted from the document.")

        # 2. Chunk and classify
        chunker = ClientDataChunker()
        chunks = chunker.process(
            text=text,
            company_id=company_id_str,
            source_file=filename,
            company_name=company_name,
        )

        # 3. Embed and index into Vector Store
        if chunks:
            try:
                from src.retrieval.bge_m3_embedder import BGEM3Embedder
                embedder = BGEM3Embedder(use_fp16=True)
                vector_store = VectorStore()

                texts = [c.enriched_text for c in chunks]
                ids = [c.chunk_id for c in chunks]
                metadatas = [c.metadata for c in chunks]

                vectors = embedder.embed_chunks(texts, batch_size=8)
                vector_store.add_chunks(
                    collection_name=COLLECTION_CLIENT,
                    ids=ids,
                    documents=texts,
                    metadatas=metadatas,
                    dense_vecs=vectors["dense"],
                    sparse_vecs=vectors["sparse"],
                )
                logger.info(f"Indexed {len(chunks)} chunks into Vector Store for '{filename}'")
            except Exception as embed_err:
                # Embedding is what makes an uploaded document usable: without
                # vectors the retrieval layer cannot see it, so every draft
                # generated afterwards silently omits its content. Marking the
                # upload "done" in that state told the user their document had
                # been processed when it had not. Fail loudly instead.
                logger.exception(f"Embedding failed for '{filename}'")
                raise RuntimeError(
                    f"Document text was extracted but could not be indexed: {embed_err}"
                ) from embed_err

        # 4. Auto-populate financials if this looks like a financial document
        if record and record.doc_type in ("financial_statement", "0"):
            _try_auto_populate_financials(text, company_id_uuid, filename, db)

        # 4b. Extract and persist structured financial tables (FinancialTableParser)
        if record and record.doc_type in ("financial_statement", "0"):
            try:
                from src.ingestion.financial_table_parser import FinancialTableParser
                from src.extraction.schema import FinancialTable as FinancialTableModel
                parser = FinancialTableParser()
                tables = parser.extract(file_path)
                for ft in tables:
                    db.add(FinancialTableModel(
                        company_id=company_id_uuid,
                        source_file=filename,
                        page=ft.page,
                        statement_type=ft.statement_type,
                        years=ft.years,
                        table_json=ft.to_dict(),
                        units=ft.units,
                        currency=ft.currency,
                    ))
                db.commit()
                logger.info(
                    f"Persisted {len(tables)} financial tables from '{filename}'"
                )
            except Exception as e:
                logger.warning(
                    f"Financial table extraction failed for '{filename}': {e}"
                )

        # 4c. Gemini structured table extraction for corporate / director documents
        if record and record.doc_type in ("board_resolution", "moa_aoa", "other"):
            try:
                from src.extraction.table_extractor import extract_and_populate_directors
                count = extract_and_populate_directors(
                    file_path, company_id_uuid, db
                )
                logger.info(
                    f"Director table upsert: {count} records from '{filename}'"
                )
            except Exception as e:
                logger.warning(
                    f"Director table extraction failed for '{filename}': {e}"
                )

        # 5. Mark done
        if record:
            record.status = "done"
            db.commit()

        logger.info(f"Document processing complete for '{filename}'")

    except Exception as e:
        logger.error(f"Background processing failed for '{filename}': {e}")
        db.rollback()
        try:
            record = db.query(UploadedDocument).filter(UploadedDocument.id == uuid.UUID(upload_id)).first()
            if record:
                record.status = "error"
                record.error_message = str(e)[:500]
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


# ─── Endpoints ────────────────────────────────────────────────────────────────

@router.post("/upload/{company_id}", response_model=UploadResponse, status_code=202)
async def upload_document(
    company_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    doc_type: str = Form(default="other"),
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """
    Accepts a document upload for the given company.
    Saves it to disk and queues a background processing job.
    Returns immediately with 202 Accepted.
    """
    # Verify authorization
    if str(current_user["company_id"]) != company_id:
        raise HTTPException(status_code=403, detail="Not authorized for this company")
    # Validate company
    try:
        company_uuid = uuid.UUID(company_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid company_id format")

    company = db.query(Company).filter(Company.id == company_uuid).first()
    if not company:
        raise HTTPException(status_code=404, detail="Company not found")

    # Validate file extension
    filename = file.filename or "upload"
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ACCEPTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Accepted: {', '.join(ACCEPTED_EXTENSIONS)}"
        )

    # Save file to disk
    dest_dir = os.path.join(CLIENT_DOCS_DIR, company_id)
    os.makedirs(dest_dir, exist_ok=True)
    dest_path = os.path.join(dest_dir, filename)

    with open(dest_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    # Auto-classify if doc_type is 'auto'
    if doc_type == "auto":
        try:
            from src.extraction.doc_classifier import DocClassifier
            classifier = DocClassifier()
            doc_type = classifier.classify_from_path(dest_path)
            logger.info(f"Auto-classified '{filename}' as '{doc_type}'")
        except Exception as e:
            logger.warning(
                f"Auto-classification failed for '{filename}': {e}. "
                "Defaulting to 'other'."
            )
            doc_type = "other"

    # Rename file to a proper descriptive name for Vector Store
    DOC_TYPE_NAMES = {
        "0": "Audited_Financial_Statements",
        "1": "Board_Resolution",
        "2": "Factory_Licence",
        "3": "Pollution_Certificate",
        "4": "Factory_Insurance",
        "5": "Trademark_Certificates",
        "6": "Vendor_Contracts",
        "7": "Litigation_Notices",
        "8": "GST_Registration",
        "9": "MOA_AOA"
    }
    
    clean_name = DOC_TYPE_NAMES.get(doc_type, "Document")
    short_id = str(uuid.uuid4())[:6]
    clean_filename = f"{clean_name}_{short_id}{ext}"
    new_dest_path = os.path.join(dest_dir, clean_filename)
    
    os.rename(dest_path, new_dest_path)
    filename = clean_filename
    dest_path = new_dest_path

    # Create upload record
    upload_record = UploadedDocument(
        company_id=company_uuid,
        filename=filename,
        doc_type=doc_type,
        status="pending",
    )
    db.add(upload_record)
    db.commit()
    db.refresh(upload_record)
    upload_id = str(upload_record.id)

    # Write audit log entry
    try:
        audit = AuditLog(
            event_type="document_upload",
            company_id=company_uuid,
            source_file=filename,
            model_used="none",
        )
        db.add(audit)
        db.commit()
        logger.info(
            f"Audit log written for upload of '{filename}' "
            f"by company {company_id}"
        )
    except Exception as e:
        logger.warning(f"Audit log write failed for '{filename}': {e}")

    # Queue background job
    background_tasks.add_task(
        _process_document_background,
        file_path=dest_path,
        upload_id=upload_id,
        company_id_str=company_id,
        company_name=company.name,
        filename=filename,
    )

    return UploadResponse(
        upload_id=upload_id,
        filename=filename,
        status="pending",
        message="File received. Processing started in the background.",
    )


@router.get("/status/{company_id}", response_model=List[DocumentStatusItem])
def get_document_status(
    company_id: str, 
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Returns the processing status of all uploaded documents for a company.
    Poll this endpoint after upload to track progress.
    """
    # Verify authorization
    if str(current_user["company_id"]) != company_id:
        raise HTTPException(status_code=403, detail="Not authorized for this company")
    try:
        company_uuid = uuid.UUID(company_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid company_id format")

    records = (
        db.query(UploadedDocument)
        .filter(UploadedDocument.company_id == company_uuid)
        .order_by(UploadedDocument.uploaded_at.desc())
        .all()
    )

    return [
        DocumentStatusItem(
            upload_id=str(r.id),
            filename=r.filename,
            doc_type=r.doc_type,
            status=r.status,
            error_message=r.error_message,
            uploaded_at=r.uploaded_at.isoformat() if r.uploaded_at else "",
        )
        for r in records
    ]


@router.delete("/{upload_id}", status_code=200)
def delete_document(
    upload_id: str,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Deletes an uploaded document, its physical file, associated vector chunks,
    and any auto-extracted data (e.g. financials).
    """
    try:
        upload_uuid = uuid.UUID(upload_id)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid upload_id format")

    record = db.query(UploadedDocument).filter(UploadedDocument.id == upload_uuid).first()
    if not record:
        raise HTTPException(status_code=404, detail="Document not found")
        
    company_id_str = str(record.company_id)
    if str(current_user["company_id"]) != company_id_str:
        raise HTTPException(status_code=403, detail="Not authorized for this company")
        
    filename = record.filename
    
    # 1. Delete chunks from Vector Store
    try:
        vector_store = VectorStore()
        deleted_count = vector_store.delete_by_metadata(
            collection_name=COLLECTION_CLIENT,
            where={"$and": [{"company_id": company_id_str}, {"source_file": filename}]}
        )
        logger.info(f"Deleted {deleted_count} chunks from Vector Store for '{filename}'")
    except Exception as e:
        logger.warning(f"Failed to delete vectors for '{filename}': {e}")
        
    # 2. Delete auto-extracted financials if this was the financial statement (index "0")
    if record.doc_type == "0":
        try:
            db.query(FinancialStatement).filter(
                FinancialStatement.company_id == record.company_id,
                FinancialStatement.source == "ai_extracted"
            ).delete()
            logger.info(f"Deleted auto-extracted financial statements for company '{company_id_str}'")
        except Exception as e:
            logger.warning(f"Failed to delete financial statements: {e}")
            
    # 3. Delete physical file
    dest_path = os.path.join(CLIENT_DOCS_DIR, company_id_str, filename)
    if os.path.exists(dest_path):
        try:
            os.remove(dest_path)
            logger.info(f"Deleted physical file '{dest_path}'")
        except OSError as e:
            logger.warning(f"Could not delete physical file '{dest_path}': {e}")
            
    # 4. Clear Generated Sections (Cascade reset)
    try:
        from src.extraction.schema import GeneratedSection
        db.query(GeneratedSection).filter(GeneratedSection.company_id == record.company_id).update(
            {"draft_text": "", "status": "pending", "completeness_score": 0.0, "flagged_gaps": []},
            synchronize_session=False
        )
        logger.info(f"Cleared generated sections for company '{company_id_str}'")
    except Exception as e:
        logger.warning(f"Failed to clear generated sections: {e}")
        
    # 5. Delete DB record
    db.delete(record)
    db.commit()
    
    return {"message": "Document and associated data successfully deleted"}
