"""
Assembles drafted sections into a DOCX/PDF that follows real SEBI filing structure.

Previously this module ordered sections by a hardcoded flat 25-name dict and
emitted `heading(name) + one paragraph of raw markdown` per section — no cover
page, no table of contents, no SECTION grouping, and markdown syntax rendered
literally. It was also dead code: nothing outside a test imported it, while the
mounted /api/export/full endpoint returned a plain-text stub.

Ordering and grouping now come from src/config/sections.py, which is derived
from the tables of contents of the 20 filings in Original_Docs/Precedents.
"""
import os
import re
import uuid
from typing import Dict, List, Tuple

from sqlalchemy.orm import Session

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

from src.config.sections import canonical_name, resolve, sort_key
from src.extraction.schema import GeneratedSection

# Statuses that represent reviewed, sign-off-ready content.
APPROVED_STATUSES = ("promoter_reviewed", "intermediary_certified")


def sort_by_sebi_toc(sections: List[GeneratedSection]) -> List[GeneratedSection]:
    """Order sections as a filed DRHP orders them; unknown names sort last."""
    return sorted(sections, key=lambda s: sort_key(s.section_name))


def group_sections(sections: List[GeneratedSection]) -> List[Tuple[str, str, List[GeneratedSection]]]:
    """
    Group ordered sections under their SECTION heading.

    Returns [(group, group_title, [section, ...]), ...]. Sections whose name
    cannot be resolved are collected under a final "UNCLASSIFIED" group rather
    than silently dropped or wedged into an arbitrary chapter.
    """
    grouped: List[Tuple[str, str, List[GeneratedSection]]] = []
    for section in sort_by_sebi_toc(sections):
        canonical = resolve(section.section_name)
        group = canonical.group if canonical else "UNCLASSIFIED"
        title = canonical.group_title if canonical else "Unclassified Sections"
        if not grouped or grouped[-1][0] != group:
            grouped.append((group, title, []))
        grouped[-1][2].append(section)
    return grouped


# ---------------------------------------------------------------------------
# Lightweight markdown rendering
# ---------------------------------------------------------------------------

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET = re.compile(r"^\s*[-*•]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[\.\)]\s+(.*)$")


def _add_markdown(doc: Document, text: str, base_level: int = 2) -> None:
    """
    Render the subset of markdown the drafting model actually emits: headings,
    bullets, numbered lists and bold spans. Previously the whole draft went in
    as a single paragraph, so '## Heading' and '**bold**' appeared literally in
    the exported document.
    """
    for raw_line in (text or "").split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue

        heading = _HEADING.match(line)
        if heading:
            level = min(base_level + len(heading.group(1)) - 1, 9)
            doc.add_heading(_strip_marks(heading.group(2)), level=level)
            continue

        bullet = _BULLET.match(line)
        if bullet:
            _add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            continue

        numbered = _NUMBERED.match(line)
        if numbered:
            _add_runs(doc.add_paragraph(style="List Number"), numbered.group(1))
            continue

        _add_runs(doc.add_paragraph(), line)


def _add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** spans."""
    position = 0
    for match in _BOLD.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def _strip_marks(text: str) -> str:
    return _BOLD.sub(r"\1", text or "").strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def build_docx(
    sections: List[GeneratedSection],
    output_path: str,
    company_name: str = "",
    include_drafts: bool = True,
) -> None:
    doc = Document()

    # Cover page
    title = doc.add_heading("DRAFT RED HERRING PROSPECTUS", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("(Subject to completion and revision)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if company_name:
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.add_run(company_name.upper()).bold = True
    if include_drafts:
        notice = doc.add_paragraph()
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = notice.add_run(
            "This export includes sections that have not been approved. "
            "It is a working draft, not a filing-ready document."
        )
        run.italic = True
    doc.add_page_break()

    groups = group_sections(sections)

    # Table of contents. Word cannot compute page numbers for a generated
    # document without a field update, so this lists structure only rather than
    # printing page numbers that would be wrong.
    doc.add_heading("TABLE OF CONTENTS", level=1)
    for group, group_title, group_sections_list in groups:
        heading = doc.add_paragraph()
        heading.add_run(f"{group} – {group_title.upper()}").bold = True
        for section in group_sections_list:
            doc.add_paragraph(canonical_name(section.section_name), style="List Bullet")
    doc.add_page_break()

    # Body
    for group, group_title, group_sections_list in groups:
        doc.add_heading(f"{group} – {group_title.upper()}", level=1)
        for section in group_sections_list:
            doc.add_heading(canonical_name(section.section_name), level=2)

            if section.status not in APPROVED_STATUSES:
                flag = doc.add_paragraph()
                flag.add_run(f"[Unapproved draft — status: {section.status}]").italic = True

            _add_markdown(doc, section.draft_text or "", base_level=3)

            if section.supporting_clause_ids:
                citation = doc.add_paragraph()
                citation.add_run("Regulatory citations: ").bold = True
                citation.add_run(", ".join(section.supporting_clause_ids))

            if getattr(section, "flagged_gaps", None):
                gaps = doc.add_paragraph()
                gaps.add_run("Outstanding gaps: ").bold = True
                gaps.add_run("; ".join(
                    g.get("description", str(g)) if isinstance(g, dict) else str(g)
                    for g in section.flagged_gaps
                ))
            doc.add_page_break()

    doc.save(output_path)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def build_pdf(
    sections: List[GeneratedSection],
    output_path: str,
    company_name: str = "",
    include_drafts: bool = True,
) -> None:
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("DRAFT RED HERRING PROSPECTUS", styles["Title"]))
    story.append(Paragraph("(Subject to completion and revision)", styles["Italic"]))
    if company_name:
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"<b>{_escape(company_name.upper())}</b>", styles["Heading2"]))
    if include_drafts:
        story.append(Spacer(1, 12))
        story.append(Paragraph(
            "<i>This export includes unapproved sections. Working draft, not a "
            "filing-ready document.</i>", styles["Normal"]))
    story.append(PageBreak())

    groups = group_sections(sections)

    story.append(Paragraph("TABLE OF CONTENTS", styles["Heading1"]))
    for group, group_title, group_sections_list in groups:
        story.append(Paragraph(f"<b>{group} – {_escape(group_title.upper())}</b>", styles["Normal"]))
        for section in group_sections_list:
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{_escape(canonical_name(section.section_name))}",
                                   styles["Normal"]))
        story.append(Spacer(1, 6))
    story.append(PageBreak())

    for group, group_title, group_sections_list in groups:
        story.append(Paragraph(f"{group} – {_escape(group_title.upper())}", styles["Heading1"]))
        for section in group_sections_list:
            story.append(Paragraph(_escape(canonical_name(section.section_name)), styles["Heading2"]))
            if section.status not in APPROVED_STATUSES:
                story.append(Paragraph(
                    f"<i>[Unapproved draft — status: {_escape(section.status)}]</i>",
                    styles["Normal"]))
            body = _escape(section.draft_text or "").replace("\n", "<br/>")
            story.append(Paragraph(body, styles["Normal"]))
            if section.supporting_clause_ids:
                story.append(Spacer(1, 6))
                story.append(Paragraph(
                    "<b>Regulatory citations:</b> " + _escape(", ".join(section.supporting_clause_ids)),
                    styles["Normal"]))
            story.append(PageBreak())

    doc.build(story)


def _escape(text: str) -> str:
    """ReportLab paragraphs parse a mini-HTML dialect, so raw &<> must be escaped."""
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def document_assembler_node(
    company_id: str,
    db: Session,
    include_drafts: bool = True,
    formats: Tuple[str, ...] = ("docx", "pdf"),
) -> Dict:
    """
    Assemble a company's drafted sections into DOCX and/or PDF.

    `include_drafts=False` restricts the export to reviewed sections — the
    filter the original code noted as missing ("In a real app we'd filter by
    status...") but never applied.
    """
    try:
        comp_uuid = uuid.UUID(str(company_id))
    except ValueError:
        raise ValueError(f"Invalid company UUID: {company_id}")

    query = db.query(GeneratedSection).filter(GeneratedSection.company_id == comp_uuid)
    if not include_drafts:
        query = query.filter(GeneratedSection.status.in_(APPROVED_STATUSES))
    sections = [s for s in query.all() if (s.draft_text or "").strip()]

    if not sections:
        detail = ("No approved sections found. Approve at least one section, or "
                  "export with drafts included.") if not include_drafts else \
                 "No drafted sections found for this company."
        return {"error": detail, "sections_included": 0}

    company = None
    try:
        from src.extraction.schema import Company
        company = db.query(Company).filter(Company.id == comp_uuid).first()
    except Exception:
        pass
    company_name = company.name if company else ""

    export_dir = os.path.join(
        os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
        "exports",
    )
    os.makedirs(export_dir, exist_ok=True)

    result = {
        "status": "success",
        "sections_included": len(sections),
        "company_name": company_name,
        "include_drafts": include_drafts,
    }

    if "docx" in formats:
        docx_path = os.path.join(export_dir, f"DRHP_{company_id}.docx")
        build_docx(sections, docx_path, company_name, include_drafts)
        result["docx_path"] = docx_path

    if "pdf" in formats:
        pdf_path = os.path.join(export_dir, f"DRHP_{company_id}.pdf")
        build_pdf(sections, pdf_path, company_name, include_drafts)
        result["pdf_path"] = pdf_path

    return result
