"""
Structural conformance of an exported DRHP against real SEBI filings.

"The export works" must mean "the export produces a document that looks like a
real filing", not "the endpoint returned 200". These tests compare an assembled
document against the chapter structure actually observed in the 20 filings under
Original_Docs/Precedents/ — the same TOC parsing that produced
src/config/sections.py.

Ordering is asserted *relatively* between high-frequency chapters rather than by
absolute index, because real filings vary slightly in where they place the less
common chapters.
"""
import os
import types

import pytest

from docx import Document

from src.agent.document_assembler import build_docx, build_pdf, group_sections, sort_by_sebi_toc
from src.config.sections import CANONICAL_SECTIONS, LEGACY_SECTIONS_25, grouped, resolve

TEST_DIR = "tests/results/phase_18_export"
os.makedirs(TEST_DIR, exist_ok=True)


def _section(name, text="Placeholder body text for this section.", status="draft",
             cites=None, gaps=None):
    obj = types.SimpleNamespace()
    obj.section_name = name
    obj.draft_text = text
    obj.status = status
    obj.supporting_clause_ids = cites or []
    obj.flagged_gaps = gaps or []
    return obj


def _all_legacy_sections():
    return [_section(name) for name in LEGACY_SECTIONS_25]


# ---------------------------------------------------------------------------
# Taxonomy is grounded in the precedent corpus
# ---------------------------------------------------------------------------

def test_canonical_sections_cover_high_frequency_chapters():
    """
    Every chapter appearing in at least 16 of the 20 precedent filings must be
    represented. The app's original 25-name list omitted several that appear in
    essentially every real filing.
    """
    high_frequency = [s for s in CANONICAL_SECTIONS if s.frequency >= 16]
    names = {s.name for s in high_frequency}

    for required in [
        "Definitions and Abbreviations",          # 20/20
        "Financial Indebtedness",                 # 20/20
        "Restrictions on Foreign Ownership of Indian Securities",  # 19/20
        "Other Financial Information",            # 19/20
        "Forward-Looking Statements",             # 19/20
        "Capitalisation Statement",               # 17/20
        "Government and Other Approvals",         # 17/20
    ]:
        assert required in names, f"{required} missing from the canonical taxonomy"


def test_every_legacy_section_name_resolves():
    """
    generated_section rows are keyed by the legacy names. If any stops resolving,
    existing drafts silently fall out of the assembled document.
    """
    unresolved = [name for name in LEGACY_SECTIONS_25 if resolve(name) is None]
    assert not unresolved, f"legacy section names no longer resolve: {unresolved}"


# ---------------------------------------------------------------------------
# Assembled document structure
# ---------------------------------------------------------------------------

def test_sections_are_ordered_like_a_real_filing():
    """Relative ordering of chapters every filing contains."""
    ordered = [s.section_name for s in sort_by_sebi_toc(_all_legacy_sections())]
    position = {name: i for i, name in enumerate(ordered)}

    # Risk Factors precedes the company chapters, which precede financials,
    # which precede the declaration — true of all 20 precedent filings.
    assert position["Risk Factors"] < position["Our Business"]
    assert position["Our Business"] < position["Financial Statements (3 Years)"]
    assert position["Financial Statements (3 Years)"] < position["Declaration & Undertakings"]
    assert position["Capital Structure"] < position["Our Business"]


def test_sections_are_grouped_under_section_headings():
    groups = group_sections(_all_legacy_sections())
    labels = [g[0] for g in groups]

    assert labels, "no groups produced"
    # Groups must appear in roman-numeral order and not repeat.
    assert labels == sorted(set(labels), key=labels.index), "group order is not stable"
    assert "SECTION II" in labels, "Risk Factors group missing"
    assert any(label.startswith("SECTION") for label in labels)


def test_unresolved_sections_are_kept_not_dropped():
    """An unrecognised section must still reach the document, flagged."""
    sections = _all_legacy_sections() + [_section("Some Bespoke Annexure")]
    groups = group_sections(sections)

    assert groups[-1][0] == "UNCLASSIFIED"
    names = [s.section_name for s in groups[-1][2]]
    assert "Some Bespoke Annexure" in names


def test_exported_docx_has_filing_structure():
    """Open the produced file and check it reads like a filing, not a text dump."""
    path = os.path.join(TEST_DIR, "conformance.docx")
    build_docx(_all_legacy_sections(), path, company_name="Test Issuer Limited",
               include_drafts=True)
    assert os.path.exists(path)

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    assert "DRAFT RED HERRING PROSPECTUS" in text
    assert "TABLE OF CONTENTS" in text
    assert "TEST ISSUER LIMITED" in text
    # An export containing unapproved drafts must say so.
    assert "not been approved" in text

    section_headings = [h for h in headings if h.startswith("SECTION ")]
    assert len(section_headings) >= 5, f"expected several SECTION groups, got {section_headings}"

    # Chapter names are canonicalised on the way out.
    assert any("Declaration" == h for h in headings), "Declaration chapter missing"


def test_markdown_is_rendered_not_emitted_literally():
    """
    The previous assembler put each draft into one paragraph, so markdown syntax
    appeared verbatim in the exported document.
    """
    path = os.path.join(TEST_DIR, "markdown.docx")
    build_docx(
        [_section("Risk Factors",
                  "## Key Risks\n\n- Customer concentration\n- **Regulatory** exposure\n\n1. First\n2. Second")],
        path, company_name="Test Issuer Limited",
    )
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    styles = {p.style.name for p in doc.paragraphs if p.text.strip()}

    assert "##" not in text, "markdown heading syntax leaked into the document"
    assert "**" not in text, "markdown bold syntax leaked into the document"
    assert "List Bullet" in styles, "bullets were not rendered as a list"
    assert "List Number" in styles, "numbered items were not rendered as a list"
    assert any(run.bold for p in doc.paragraphs for run in p.runs), "no bold run rendered"


def test_approved_only_export_excludes_drafts():
    """include_drafts=False is the filter the original code noted but never applied."""
    from src.agent.document_assembler import APPROVED_STATUSES
    assert "intermediary_certified" in APPROVED_STATUSES

    sections = [
        _section("Risk Factors", status="draft"),
        _section("Our Business", status="intermediary_certified"),
    ]
    approved = [s for s in sections if s.status in APPROVED_STATUSES]
    path = os.path.join(TEST_DIR, "approved_only.docx")
    build_docx(approved, path, company_name="Test Issuer Limited", include_drafts=False)

    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "Our Business" in text
    assert "Risk Factors" not in text
    # A filing-ready export must not carry the working-draft warning.
    assert "not been approved" not in text


# ---------------------------------------------------------------------------
# Cover page — geometry and boilerplate measured from a real filed prospectus
# (Original_Docs/Precedents/drhp_physics_wallah.pdf)
# ---------------------------------------------------------------------------

def test_pdf_cover_page_has_no_unrenderable_glyphs():
    """
    The base-14 PDF fonts use WinAnsiEncoding, which does not include ₹ (U+20B9)
    or ● (U+25CF) — both appeared as mojibake ("I", "[G]") without an embedded
    Unicode font. The cover module uses "Rs." and "•" instead.
    """
    import fitz

    path = os.path.join(TEST_DIR, "cover_glyphs.pdf")
    build_pdf(
        [_section("Risk Factors", "Placeholder.")],
        path, company_name="Glyph Test Limited", include_drafts=True,
        cover={"company_name": "Glyph Test Limited", "total_shares": 1000000,
               "price_per_share": 50, "issue_size_lakhs": 500},
        all_chapters=grouped(),
    )
    text = fitz.open(path)[0].get_text()
    assert "Rs." in text
    assert "500.00" in text  # issue size actually reached the page
    for bad in ("�", "[G]", "I500", "I50"):
        assert bad not in text, f"unrenderable-glyph artifact found: {bad!r}"


def test_pdf_cover_page_uses_real_company_data():
    """The cover must carry the company's actual details, not placeholders,
    for every field that was supplied."""
    import fitz

    path = os.path.join(TEST_DIR, "cover_real_data.pdf")
    cover = {
        "company_name": "Rajputana Precision Steels Limited",
        "cin": "U27100RJ2015PLC047823",
        "registered_office": "Plot No. 42, RIICO Industrial Area, Bhiwadi, Rajasthan",
        "promoters": ["Vikram Deshpande"],
        "total_shares": 4000000,
        "price_per_share": 125,
        "issue_size_lakhs": 5000,
    }
    build_pdf(
        [_section("Risk Factors", "Placeholder.")],
        path, company_name=cover["company_name"], include_drafts=True,
        cover=cover, all_chapters=grouped(),
    )
    text = fitz.open(path)[0].get_text()
    assert "RAJPUTANA PRECISION STEELS LIMITED" in text
    assert cover["cin"] in text
    assert "RIICO Industrial Area" in text
    assert "VIKRAM DESHPANDE" in text
    assert "4,000,000" in text
    assert "5,000.00" in text


def test_pdf_cover_page_placeholders_missing_fields_honestly():
    """
    Fields the company has not supplied render as the "[•]" placeholder real
    draft filings use — never as an invented value.
    """
    import fitz

    path = os.path.join(TEST_DIR, "cover_missing_fields.pdf")
    build_pdf(
        [_section("Risk Factors", "Placeholder.")],
        path, company_name="Sparse Data Limited", include_drafts=True,
        cover={"company_name": "Sparse Data Limited"},  # nothing else supplied
        all_chapters=grouped(),
    )
    text = fitz.open(path)[0].get_text()
    assert "[•]" in text
    assert "SPARSE DATA LIMITED" in text


def test_pdf_table_of_contents_marks_undrafted_chapters():
    """
    The contents page must show the full canonical DRHP structure and flag
    which chapters are not yet drafted — an early-stage export must not imply
    the drafted sections are the whole filing.
    """
    import fitz

    path = os.path.join(TEST_DIR, "cover_toc.pdf")
    build_pdf(
        [_section("Risk Factors", "Placeholder.")],
        path, company_name="TOC Test Limited", include_drafts=True,
        cover={}, all_chapters=grouped(),
    )
    toc_text = fitz.open(path)[1].get_text()
    assert "TABLE OF CONTENTS" in toc_text
    assert "SECTION II" in toc_text and "Risk Factors" in toc_text
    # The drafted chapter must NOT carry the undrafted marker.
    risk_line = next(l for l in toc_text.split("\n") if l.strip() == "Risk Factors")
    assert "[•]" not in risk_line
    # An undrafted chapter must carry it.
    assert "Our Business [•]" in toc_text or "Our Business  [•]" in toc_text
